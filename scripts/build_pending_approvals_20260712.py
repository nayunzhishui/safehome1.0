"""Build the remaining-approval register after task 18 owner approvals."""

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from build_task18_mentor_report import BLUE, add_bullets, add_heading, add_kv_table, add_note, add_table


OUTPUT = Path(r"D:\桌面\Desktop\待批准内容说明_20260712.docx")


def configure(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.25
    normal.paragraph_format.space_after = Pt(6)
    for name, size in (("Title", 24), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.paragraph_format.keep_with_next = True
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "安心家任务十八  |  待批准与外部验收事项"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(107, 114, 128)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    footer._p.append(field)
    footer.add_run(" 页")


def main() -> None:
    doc = Document()
    configure(doc)

    title = doc.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("安心家项目待批准内容说明")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("任务十八完成后的治理、伦理、外部效度与真实环境验收清单\n更新日期：2026年7月12日")

    add_note(
        doc,
        "本文件不重复要求项目负责人批准已经批准的量表、画像模型和三个项目试点。项目负责人批准已记录为项目内部准入决定；"
        "它不能替代量表版权许可、正式中文版来源核验、独立心理学评审、研究伦理审查、真实部署验收或模型外部效度证据。",
    )

    add_heading(doc, "一、已经完成的项目负责人批准", 1)
    add_kv_table(
        doc,
        [
            ("批准范围", "33 份可执行支持性测评、当前画像模型及三个项目试点的项目内部准入"),
            ("系统状态", "33 份 worksheet 已设为 pilot_approved 并开放；11 个画像模型中 10 个当前模型开放，1 个历史 HPLP 模型仅保留比较"),
            ("批准性质", "项目负责人关于本项目试点使用与功能开放的内部决定"),
            ("批准记录", r"D:\codex\workspace\safehome1.0\docs\02_专项进度与验收\任务十八项目负责人量表与画像试点批准记录_20260712.md"),
            ("自动验收", "量表内容审计 33 份，blocker/error/warning 均为 0；后端 278 项测试通过；Web 构建、类型检查及 12 项端到端测试通过"),
        ],
    )

    add_heading(doc, "二、仍需批准或补充证据的事项", 1)
    rows = [
        ["A1", "量表版权与正式版本来源", "确认每份量表的授权范围、正式中文版来源、题项转载与数字化施测权限；公开可用不等于可任意转载。", "资料/版权负责人", "授权文件、原作者声明、正式出版物或明确开放许可可追溯；版本号与题项一致", "正式扩大使用前"],
        ["A2", "心理测量内容独立复核", "对题项、选项锚点、反向计分、维度归属、总分/均分算法、缺失值规则和解释边界进行独立复核。", "心理学专家", "逐量表签字；抽测题项分值可复算；不把支持性结果写成诊断结论", "试点正式收数前"],
        ["A3", "三个项目的研究与伦理准入", "复核招募对象、知情同意、未成年人监护人同意、退出机制、风险升级、数据用途和保留期限。", "伦理/研究负责人", "伦理意见或会议纪要；知情同意版本号；风险联系人和转介流程明确", "招募开放前"],
        ["A4", "画像模型外部效度", "对当前聚类命名、K 值、稳定性、低置信度阈值、离群处理和训练卡映射进行独立数据复核。", "统计/心理测量专家", "留出样本或新样本复核；稳定性和解释一致性达到预设标准；保留不通过阻断", "画像用于正式反馈前"],
        ["A5", "情感计算标注与误差", "验证情绪词典、否定/程度处理、上下文歧义、低置信度处理及人工复核策略。", "内容专家+数据分析人员", "双人标注样本、分歧记录、准确率/一致性报告、误用案例清单", "真实文本批量分析前"],
        ["A6", "社会网络与家庭拓扑解释边界", "确认关系节点、边、方向、权重和时间窗口不会被误读为家庭诊断或价值判断。", "心理学/家庭研究专家", "示例案例复核；解释文案签字；高风险关系数据仅供授权研究者查看", "面向用户展示前"],
        ["A7", "CloudBase 正式部署", "复核生产环境变量、MySQL 迁移、备份恢复、网络访问、域名、日志脱敏和服务监控。", "部署/后端负责人", "生产健康检查、迁移记录、回滚演练、备份恢复证明、无明文密钥", "正式发布前"],
        ["A8", "真实登录与角色权限", "在真实微信、手机号、学生、家长和研究者账号上验证身份衔接、越权保护和退出登录。", "产品负责人+测试人员", "真机验收记录；角色矩阵逐项通过；匿名数据合并无串号", "试点账号发放前"],
        ["A9", "多端视觉与可用性", "在 Android、iOS、微信开发者工具和 Web 常用视口检查表格、Canvas、雷达图、长图和交互。", "UI/测试人员", "关键页面截图；无重叠、裁切、不可点击；无障碍关键项通过", "版本提审前"],
        ["A10", "数据保留、撤回与删除", "确认数据保存期限、撤回后处理、删除请求、研究留档例外和脱敏导出口径。", "伦理/数据治理负责人", "正式政策文本、责任人、工单流程和审计记录要求明确", "正式收集数据前"],
    ]
    add_table(doc, ["编号", "事项", "需要确认的内容", "建议批准人", "通过证据", "最迟节点"], rows, [620, 1320, 2850, 1350, 2400, 1050])

    add_heading(doc, "三、建议优先级与执行顺序", 1)
    add_bullets(
        doc,
        [
            "P0：A2、A3、A8、A10。它们直接决定是否可以安全招募、收数和按角色开放数据。",
            "P1：A1、A4、A7、A9。它们决定量表长期使用、画像可信度和生产环境可用性。",
            "P2：A5、A6。当前离线原型可继续研究，但在完成标注与解释边界验收前，不应作为自动化结论直接推送给用户。",
        ],
    )

    add_heading(doc, "四、批准时应填写的最小信息", 1)
    add_table(
        doc,
        ["字段", "填写要求"],
        [
            ["事项编号", "使用本清单 A1-A10 编号"],
            ["决定", "批准 / 附条件批准 / 退回修改"],
            ["范围", "明确适用量表、模型、项目、角色、环境和版本"],
            ["证据位置", "填写文件绝对路径、会议纪要编号或正式链接"],
            ["限制条件", "试点人数、有效期、仅研究者可见、不得导出等"],
            ["批准人和日期", "保留姓名、角色、日期和版本；不要把账号密码写入本文件"],
        ],
        [1900, 7440],
    )

    add_heading(doc, "五、系统继续保留的阻断规则", 1)
    add_bullets(
        doc,
        [
            "低置信度、离群、数据不足、高风险和未批准模型继续进入阻断或人工复核状态。",
            "高风险文本不生成普通自动训练建议，转入人工支持或既定转介流程。",
            "缺少可执行 worksheet 的目录项不会因为项目负责人批准而自动开放。",
            "历史 HPLP 模型只用于方法比较；当前健康生活方式画像使用牛至旭研究所对应的 40 题、六维 HPLP 版本。",
            "所有反馈使用阶段性观察、自我了解线索和可练习建议，不输出诊断、人格定性或治疗承诺。",
        ],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
