from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


REPO_ROOT = Path(__file__).resolve().parents[1]
DRAFTS_PATH = REPO_ROOT / "content" / "scale_item_drafts.json"
WORKSHEETS_PATH = REPO_ROOT / "content" / "assessment_worksheets.json"
OUTPUT_PATH = Path(r"D:\codex\workspace\safehome1.0其他内容\量表题项内容人工审核对照.docx")

SCALE_ID_ALIASES = {
    "emotion_regulation_erq_gross": "emotion_regulation_erq",
}


def load_json(path: Path) -> dict:
    with path.open("r", encoding="utf-8") as f:
        return json.load(f)


def normalize_text(value: object) -> str:
    if value is None:
        return ""
    text = str(value)
    text = re.sub(r"\s+", "", text)
    text = text.replace("（", "(").replace("）", ")")
    text = text.replace("，", ",").replace("。", ".")
    text = text.replace("：", ":").replace("；", ";")
    return text.strip()


def set_run_font(run, size: float | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_cell_text(cell, text: object, bold: bool = False, size: float = 8.5) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.space_after = Pt(0)
    run = paragraph.add_run("" if text is None else str(text))
    set_run_font(run, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_inches: float) -> None:
    cell.width = Inches(width_inches)
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(int(width_inches * 1440)))
    tc_w.set(qn("w:type"), "dxa")


def set_table_widths(table, widths: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths):
            if idx < len(row.cells):
                set_cell_width(row.cells[idx], width)


def add_table(doc: Document, headers: list[str], rows: list[list[object]], widths: list[float], font_size: float = 8.0):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    set_table_widths(table, widths)
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, size=font_size)
        shade_cell(table.rows[0].cells[idx], "E8EEF5")
    for row_values in rows:
        row = table.add_row()
        for idx, value in enumerate(row_values):
            set_cell_text(row.cells[idx], value, size=font_size)
    set_table_widths(table, widths)
    return table


def add_heading(doc: Document, text: str, level: int) -> None:
    paragraph = doc.add_heading("", level=level)
    paragraph.paragraph_format.space_before = Pt(8 if level == 1 else 5)
    paragraph.paragraph_format.space_after = Pt(4)
    run = paragraph.add_run(text)
    set_run_font(run, size={1: 15, 2: 12, 3: 10.5}.get(level, 10), bold=True, color="1F4D78")


def add_paragraph(doc: Document, text: str, size: float = 9.5, bold: bool = False) -> None:
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(text)
    set_run_font(run, size=size, bold=bold)


def options_text(options: list[dict]) -> str:
    if not options:
        return "无"
    parts = []
    for option in options:
        value = option.get("value", "")
        score = option.get("score", value)
        label = option.get("label", "")
        parts.append(f"{value}/{score}: {label}")
    return "；".join(parts)


def draft_options_text(draft: dict) -> str:
    options = draft.get("likert") or draft.get("options") or []
    if not options:
        return "无"
    parts = []
    for option in options:
        parts.append(f"{option.get('value', '')}: {option.get('label', '')}")
    return "；".join(parts)


def worksheet_options_text(worksheet: dict | None) -> str:
    if not worksheet:
        return "未进入小程序题库"
    questions = worksheet.get("questions") or []
    first_scale_question = next((q for q in questions if q.get("options")), None)
    if not first_scale_question:
        return "无"
    return options_text(first_scale_question.get("options") or [])


def item_match_status(draft_item: dict, worksheet_question: dict | None) -> str:
    if not worksheet_question:
        return "小程序缺失"
    draft_text = normalize_text(draft_item.get("text"))
    worksheet_text = normalize_text(worksheet_question.get("prompt"))
    return "一致" if draft_text == worksheet_text else "不一致"


def build_review_rows(draft: dict, worksheet: dict | None) -> tuple[list[list[object]], dict]:
    questions = worksheet.get("questions") if worksheet else []
    by_id = {question.get("id"): question for question in questions}
    by_order = {idx + 1: question for idx, question in enumerate(questions)}
    rows: list[list[object]] = []
    status_counts = {"一致": 0, "不一致": 0, "小程序缺失": 0}

    for item in sorted(draft.get("items") or [], key=lambda x: x.get("display_order", 0)):
        order = item.get("display_order", "")
        question = by_id.get(item.get("item_code")) or by_order.get(order)
        status = item_match_status(item, question)
        status_counts[status] = status_counts.get(status, 0) + 1
        rows.append(
            [
                order,
                item.get("item_code", ""),
                item.get("text", ""),
                question.get("prompt", "") if question else "",
                item.get("dimension", ""),
                "是" if item.get("reverse_scored") else "否",
                status,
            ]
        )
    return rows, status_counts


def worksheet_for_draft(draft: dict, worksheets_by_id: dict[str, dict]) -> tuple[str, dict | None]:
    draft_id = draft.get("scale_id", "")
    worksheet_id = SCALE_ID_ALIASES.get(draft_id, draft_id)
    return worksheet_id, worksheets_by_id.get(worksheet_id)


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(11)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.55)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(9.5)


def main() -> None:
    drafts_data = load_json(DRAFTS_PATH)
    worksheets_data = load_json(WORKSHEETS_PATH)
    drafts = [d for d in drafts_data.get("drafts", []) if d.get("items")]
    worksheets_by_id = {w.get("id"): w for w in worksheets_data.get("worksheets", [])}

    doc = Document()
    setup_document(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("量表题项内容人工审核对照")
    set_run_font(run, size=20, bold=True, color="000000")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(8)
    subtitle_run = subtitle.add_run(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M')} | 来源：scale_item_drafts.json 与 assessment_worksheets.json")
    set_run_font(subtitle_run, size=9, color="555555")

    add_paragraph(
        doc,
        "用途：本文件用于人工核对已录入草稿量表的题项内容、题序、选项、维度和反向题信息。"
        "本文件不包含被试逐行原始数据，也不代表量表已经完成授权、计分和用户端开放审核。",
        size=9,
    )

    overview_rows = []
    per_scale_data = []
    for idx, draft in enumerate(drafts, start=1):
        worksheet_id, worksheet = worksheet_for_draft(draft, worksheets_by_id)
        item_rows, status_counts = build_review_rows(draft, worksheet)
        draft_count = len(draft.get("items") or [])
        worksheet_count = len(worksheet.get("questions") or []) if worksheet else 0
        quantity_status = "一致" if draft_count == worksheet_count else "不一致"
        text_status = "一致" if status_counts.get("不一致", 0) == 0 and status_counts.get("小程序缺失", 0) == 0 else "需复核"
        focus = []
        if quantity_status != "一致":
            focus.append("题项数量")
        if text_status != "一致":
            focus.append("题目文字")
        if draft.get("scoring_status") not in {"rule_available_pending_review", "draft_from_syntax_pending_review", "draft_from_pdf_and_sps_pending_review"}:
            focus.append("计分规则")
        if not focus:
            focus.append("常规复核")

        overview_rows.append(
            [
                idx,
                draft.get("scale_id", ""),
                worksheet_id,
                draft.get("display_name", ""),
                draft_count,
                worksheet_count,
                quantity_status,
                text_status,
                "；".join(focus),
            ]
        )
        per_scale_data.append((idx, draft, worksheet_id, worksheet, item_rows, status_counts))

    add_heading(doc, "1. 总览", 1)
    add_table(
        doc,
        ["序号", "草稿量表ID", "小程序量表ID", "量表名称", "草稿题数", "小程序题数", "数量", "题目", "重点复核"],
        overview_rows,
        [0.35, 1.65, 1.55, 1.8, 0.55, 0.65, 0.5, 0.5, 1.35],
        font_size=7.5,
    )

    add_heading(doc, "2. 逐量表题项对照", 1)
    for idx, draft, worksheet_id, worksheet, item_rows, status_counts in per_scale_data:
        add_heading(doc, f"{idx}. {draft.get('display_name', '')}", 2)
        metadata_rows = [
            ["草稿量表ID", draft.get("scale_id", "")],
            ["小程序量表ID", worksheet_id],
            ["来源文件", "；".join(draft.get("source_files") or [])],
            ["来源目录", draft.get("source_folder", "")],
            ["用户分类", (worksheet or {}).get("category", draft.get("audience", ""))],
            ["草稿题项数", len(draft.get("items") or [])],
            ["小程序题项数", len((worksheet or {}).get("questions") or [])],
            ["计分状态", draft.get("scoring_status", "")],
            ["审查状态", (worksheet or {}).get("review_status", draft.get("review_status", ""))],
            ["题目一致性统计", f"一致 {status_counts.get('一致', 0)}；不一致 {status_counts.get('不一致', 0)}；小程序缺失 {status_counts.get('小程序缺失', 0)}"],
        ]
        add_table(doc, ["字段", "内容"], metadata_rows, [1.2, 8.6], font_size=8)

        add_heading(doc, "填写说明与选项", 3)
        add_table(
            doc,
            ["项目", "内容"],
            [
                ["草稿填写说明", draft.get("instructions", "")],
                ["草稿选项", draft_options_text(draft)],
                ["小程序首题选项", worksheet_options_text(worksheet)],
            ],
            [1.2, 8.6],
            font_size=8,
        )

        dimensions = draft.get("dimensions") or []
        if dimensions:
            add_heading(doc, "维度与反向题", 3)
            dimension_rows = []
            for dimension in dimensions:
                dimension_rows.append(
                    [
                        dimension.get("code", ""),
                        dimension.get("label", ""),
                        "、".join(dimension.get("item_codes") or dimension.get("item_ids") or []),
                        "、".join(dimension.get("reverse_item_codes") or []),
                        dimension.get("note", "") or dimension.get("description", ""),
                    ]
                )
            add_table(doc, ["维度代码", "维度名称", "题项", "反向题", "备注"], dimension_rows, [0.9, 1.4, 3.3, 1.1, 3.1], font_size=7.5)

        add_heading(doc, "题项逐条对照", 3)
        add_table(
            doc,
            ["题序", "题项ID", "草稿题目", "小程序题目", "维度", "反向", "一致性"],
            item_rows,
            [0.4, 0.8, 3.05, 3.05, 1.0, 0.45, 0.65],
            font_size=7.2,
        )

        scoring_notes = draft.get("scoring_notes") or []
        if scoring_notes:
            add_heading(doc, "计分与复核备注", 3)
            for note in scoring_notes:
                add_paragraph(doc, f"- {note}", size=8)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
