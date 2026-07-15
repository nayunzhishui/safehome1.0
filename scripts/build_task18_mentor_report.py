"""Build the 2026-07-11 mentor audit report from reviewed project content."""

from __future__ import annotations

import json
import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path(r"D:\桌面\Desktop\安心家量表画像训练与系统全链路审核汇报_20260711.docx")
BLUE = "24476B"
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F3F4F6"
WHITE = "FFFFFF"
PAGE_BREAK_PENDING = False


def load_json(relative: str) -> Any:
    return json.loads((ROOT / relative).read_text(encoding="utf-8"))


def txt(value: Any, default: str = "未配置") -> str:
    if value is None or value == "":
        return default
    if isinstance(value, bool):
        return "是" if value else "否"
    if isinstance(value, list):
        return "；".join(txt(item, "") for item in value if item not in (None, "")) or default
    if isinstance(value, dict):
        return "；".join(f"{k}：{txt(v, '')}" for k, v in value.items()) or default
    return str(value)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, twips: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(twips))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths: list[int]) -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths[min(idx, len(widths) - 1)])
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tc_pr = cell._tc.get_or_add_tcPr()
            mar = OxmlElement("w:tcMar")
            for side, value in (("top", 80), ("bottom", 80), ("left", 120), ("right", 120)):
                node = OxmlElement(f"w:{side}")
                node.set(qn("w:w"), str(value))
                node.set(qn("w:type"), "dxa")
                mar.append(node)
            tc_pr.append(mar)


def repeat_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_repeat_table_header(row) -> None:
    repeat_header(row)
    for cell in row.cells:
        set_cell_shading(cell, BLUE)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(255, 255, 255)
                run.font.bold = True


def add_table(doc: Document, headers: list[str], rows: Iterable[list[Any]], widths: list[int]) -> Any:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row_idx, values in enumerate(rows):
        cells = table.add_row().cells
        for idx, value in enumerate(values):
            cells[idx].text = txt(value, "—")
            if row_idx % 2:
                set_cell_shading(cells[idx], LIGHT_GRAY)
    set_table_geometry(table, widths)
    set_repeat_table_header(table.rows[0])
    doc.add_paragraph()
    return table


def add_kv_table(doc: Document, pairs: Iterable[tuple[str, Any]]) -> Any:
    rows = [[key, value] for key, value in pairs]
    table = add_table(doc, ["项目", "审核内容"], rows, [1850, 7510])
    for row in table.rows[1:]:
        set_cell_shading(row.cells[0], LIGHT_BLUE)
        row.cells[0].paragraphs[0].runs[0].font.bold = True
    return table


def add_bullets(doc: Document, items: Iterable[Any], numbered: bool = False) -> None:
    style = "List Number" if numbered else "List Bullet"
    for item in items:
        if item in (None, "", []):
            continue
        doc.add_paragraph(txt(item), style=style)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.3)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.italic = True
    run.font.color.rgb = RGBColor(75, 85, 99)
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), LIGHT_BLUE)
    p_pr.append(shd)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    global PAGE_BREAK_PENDING
    paragraph = doc.add_heading(text, level=level)
    if PAGE_BREAK_PENDING:
        paragraph.paragraph_format.page_break_before = True
        PAGE_BREAK_PENDING = False


def add_page_break(doc: Document) -> None:
    global PAGE_BREAK_PENDING
    PAGE_BREAK_PENDING = True


def feature_summary(model: dict[str, Any]) -> str:
    rows = []
    for feature in model.get("features", []):
        if not isinstance(feature, dict):
            continue
        feature_id = feature.get("feature_id") or feature.get("source_variable")
        worksheet_id = feature.get("worksheet_question_id")
        source = feature.get("source_variable")
        parts = [txt(feature_id, "")]
        if worksheet_id and worksheet_id != feature_id:
            parts.append(f"worksheet={worksheet_id}")
        if source and source != feature_id:
            parts.append(f"source={source}")
        rows.append("/".join(parts))
    return f"n={model.get('n_cases')}；特征数={model.get('n_features')}；" + "，".join(rows)


def dim_summary(worksheet: dict[str, Any]) -> str:
    dimensions = worksheet.get("dimensions", [])
    if not dimensions:
        return "仅总分或未设维度"
    return "；".join(
        f"{d.get('label') or d.get('code') or d.get('id')}（{len(d.get('item_ids', [])) or len(d.get('question_ids', [])) or len(d.get('terms', []))}项）"
        for d in dimensions
    )


def scoring_summary(worksheet: dict[str, Any]) -> str:
    scoring = worksheet.get("scoring", {})
    if not isinstance(scoring, dict):
        return txt(scoring)
    parts = [txt(scoring.get("method") or scoring.get("type"), "按题项分值汇总")]
    if scoring.get("reverse_question_ids"):
        parts.append(f"反向题：{txt(scoring['reverse_question_ids'])}")
    if scoring.get("total_range"):
        parts.append(f"总分范围：{txt(scoring['total_range'])}")
    if scoring.get("interpretation"):
        parts.append(f"解释：{txt(scoring['interpretation'])}")
    return "；".join(parts)


def condition_text(condition: Any) -> str:
    if not isinstance(condition, dict):
        return txt(condition)
    labels = {
        "worksheet_id": "量表",
        "dimension_id": "维度",
        "operator": "运算符",
        "value": "阈值",
        "min_score": "最低分",
        "max_score": "最高分",
        "risk_level": "风险级别",
    }
    return "；".join(f"{labels.get(k, k)}={txt(v)}" for k, v in condition.items())


def model_method(model: dict[str, Any]) -> str:
    selection = model.get("model_selection", {})
    if isinstance(selection, dict) and selection.get("selected_method"):
        return txt(selection["selected_method"])
    if model.get("mixture_weights") or model.get("diag_covariances"):
        return "GaussianMixture"
    return "KMeans"


def model_selection_summary(selection: Any) -> str:
    if isinstance(selection, dict):
        return "；".join(
            f"{k}={txt(v)}"
            for k, v in selection.items()
            if k in {"selected_method", "selection_reason", "candidate_ks", "best_k", "metric", "random_state"}
        ) or "已在冻结工件中保存候选模型指标"
    if isinstance(selection, list):
        summaries = []
        for row in selection:
            if not isinstance(row, dict):
                continue
            keys = [k for k in ("method", "k", "silhouette", "calinski_harabasz", "davies_bouldin", "bic", "aic") if k in row]
            summaries.append("，".join(f"{k}={row[k]}" for k in keys))
        return "；".join(summaries[:12]) or "已在冻结工件中保存候选模型指标"
    return txt(selection)


def add_cover(doc: Document, metrics: dict[str, int]) -> None:
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("安心家量表、聚类画像、训练内容\n与系统全链路审核汇报")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor.from_string(BLUE)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("导师审核文字材料｜更新至2026年7月12日")
    r2.font.size = Pt(14)
    r2.font.color.rgb = RGBColor(75, 85, 99)
    doc.add_paragraph()
    add_table(
        doc,
        ["内容库", "当前数量", "审核口径"],
        [
            ["小程序 worksheet", metrics["worksheets"], f"已开放 {metrics['enabled']} 个"],
            ["量表目录", metrics["catalog"], "逐项区分已开放、内部审查和未找到"],
            ["聚类画像模型", metrics["models"], "仅批准且通过门禁的模型可解释"],
            ["训练卡", metrics["cards"], "逐卡列规则、适用条件与停止规则"],
            ["项目试点方案", metrics["programs"], "均须独立审批后进入试点"],
        ],
        [2900, 1700, 4760],
    )
    add_note(doc, "用途：导师审核、试点前内容确认与研究方法复核。本文不构成诊断、治疗建议、危机预测或模型有效性证明。")
    add_page_break(doc)


def build() -> Path:
    worksheets_data = load_json("content/assessment_worksheets.json")
    worksheets = worksheets_data["worksheets"]
    catalog = load_json("content/scales_catalog.json")["scales"]
    cards = load_json("content/training_cards.json")["cards"]
    mapping = load_json("content/assessment_training_map.json")
    programs = load_json("content/programs.json")["programs"]
    courses = load_json("content/courses.json")["courses"]
    profile_paths = sorted((ROOT / "content/profiles").glob("*.json"))
    profiles = [json.loads(path.read_text(encoding="utf-8")) for path in profile_paths]
    worksheet_by_id = {w["id"]: w for w in worksheets}
    profile_by_ws: dict[str, list[dict[str, Any]]] = {}
    for model in profiles:
        linked_id = model.get("worksheet_id") or model.get("scale_id")
        if linked_id:
            profile_by_ws.setdefault(linked_id, []).append(model)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25
    for name, size in (("Title", 25), ("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11.5)):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(BLUE)
        style.font.bold = True
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.space_before = Pt(12 if name == "Heading 1" else 8)
        style.paragraph_format.space_after = Pt(6)
    for name in ("List Bullet", "List Number"):
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)

    header = section.header.paragraphs[0]
    header.text = "安心家导师审核材料  |  量表 · 画像 · 训练 · 研究治理"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(107, 114, 128)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.add_run("第 ")
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), "PAGE")
    footer._p.append(fld)
    footer.add_run(" 页")

    metrics = {
        "worksheets": len(worksheets),
        "enabled": sum(bool(w.get("enabled_for_user")) for w in worksheets),
        "catalog": len(catalog),
        "models": len(profiles),
        "cards": len(cards),
        "programs": len(programs),
    }
    add_cover(doc, metrics)

    add_heading(doc, "材料结构与审核结论", 1)
    add_bullets(
        doc,
        [
            "第一部分说明量表开放结论、题项、计分、维度、来源与未解决问题。",
            "第二部分说明聚类画像的数据、预处理、算法、画像结果和五类阻断状态。",
            "第三部分逐卡汇总训练卡内容、来源量表、推荐规则、最低剂量和停止规则。",
            "第四部分汇总三个项目试点、情感计算、语义共现网络和家庭关系拓扑。",
            "第五部分说明账户、鉴权、数据库、角色分工以及一次完整数据流。",
            "附录给出全部 worksheet 逐题清单、课程内容与证据路径，便于逐项签字。",
        ],
        numbered=True,
    )
    add_note(doc, "项目负责人已批准现有量表和画像进入试点。具备完整 worksheet、计分和维度链路的33份内容已开放；仅有目录元数据而缺少执行链的量表继续阻断。该批准不冒充第三方版权许可或独立伦理批件。")

    add_heading(doc, "一、量表导入、开放与公开查找审核", 1)
    add_heading(doc, "1.1 当前总体情况", 2)
    add_kv_table(
        doc,
        [
            ("小程序 worksheet 总数", len(worksheets)),
            ("当前面向用户开放", sum(bool(w.get("enabled_for_user")) for w in worksheets)),
            ("继续隐藏", sum(not bool(w.get("enabled_for_user")) for w in worksheets)),
            ("量表目录记录", len(catalog)),
            ("服务端计分原则", "客户端只提交题号和选择；服务端校验题号、重复、必答和选项后自行计分，忽略客户端伪造分值。"),
            ("解释边界", "结果只能作为支持性测评和阶段性观察线索，不直接生成诊断、人格定性或治疗承诺。"),
        ],
    )

    add_heading(doc, "1.2 全部导入量表总览", 2)
    overview_rows = []
    for w in worksheets:
        models = profile_by_ws.get(w["id"], [])
        overview_rows.append(
            [
                w["id"],
                w.get("display_title"),
                len(w.get("questions", [])),
                dim_summary(w),
                "已开放" if w.get("enabled_for_user") else "隐藏",
                "；".join(m["model_id"] for m in models) if models else "无直接绑定",
            ]
        )
    add_table(doc, ["worksheet_id", "显示名称", "题数", "维度", "状态", "聚类模型"], overview_rows, [1750, 1950, 650, 2550, 750, 1710])

    add_heading(doc, "1.3 重点量表的开放与版本决策", 2)
    opened_ids = [
        "parental_burnout_pba",
        "rsca_adolescent_resilience",
        "regulatory_focus_general_18",
        "rfq8_reflective_functioning",
        "self_compassion_scs_cn",
        "fmi_12_mindfulness",
        "swls_life_satisfaction",
        "perceived_social_support_psss",
        "big_five_bfi_60",
        "attribution_style_student_36",
        "ghq12_general_health",
        "micro_ysq_relationship_18",
        "big_five_tipi_10",
        "hplp_c_health_promoting_lifestyle",
    ]
    rows = []
    for wid in opened_ids:
        w = worksheet_by_id[wid]
        rows.append([wid, w.get("display_title"), len(w["questions"]), dim_summary(w), scoring_summary(w), "已开放"])
    add_table(doc, ["ID", "量表", "题数", "维度", "计分", "决定"], rows, [1750, 1500, 600, 2200, 2710, 600])
    add_table(doc, ["量表", "本轮公开核对来源", "用途"], [
        ["PBA", "https://pmc.ncbi.nlm.nih.gov/articles/PMC5998056/", "23题原始结构、四因素与父母倦怠构念边界"],
        ["RSCA", "https://journal.psych.ac.cn/xlxb/CN/abstract/abstract206.shtml", "27题青少年心理韧性量表的五维结构与中国样本来源"],
        ["一般调节焦点18题", "https://sjdm.org/dmidi/General_Regulatory_Focus_Measure.html", "Lockwood一般调节焦点量表的促进/预防双维结构"],
        ["RFQ-8", "https://peerj.com/articles/5756/", "RFQc/RFQu六题映射、非线性计分表与均值口径"],
    ], [1700, 4250, 3410])
    add_note(doc, "HPLP版本冲突已按项目负责人指定解决：运行口径统一采用牛至旭资料中的40题版本，按营养、健康责任、身体活动、精神成长、人际关系和压力管理六维取均值；原42题草稿退出运行，高鸣聪HPLP模型仅作历史对照。")

    add_heading(doc, "1.4 原4.2“尚无可靠 worksheet”的公开查找结果", 2)
    public_rows = [
        ["WHO-5", "已查到并导入", "5", "单一幸福感维度；0–5分；原始分0–25，乘4为0–100", "WHO官方2024版；CC BY-NC-SA 3.0 IGO", "https://www.who.int/publications/m/item/WHO-UCN-MSD-MHE-2024.01"],
        ["认知好奇量表 ECS-C", "已查到并导入", "10", "兴趣型1/2/3/4/5/7；剥夺型6/8/9/10；1–4分求和", "南京大学公开中文表；与用户提供SPSS口径一致", "https://rlrw.nju.edu.cn/_upload/article/files/a1/c0/bd0ca56c40be8a8e5fa98241b588/fbbad131-99dd-4acf-98f1-bc8d298df3f6.pdf"],
        ["十项目大五人格 TIPI-C", "已查到并导入", "10", "2/4/6/8/10反向；五维各2题取均值", "作者公开页面允许使用；中文译本由作者页提供", "https://gosling.psy.utexas.edu/scales-weve-developed/ten-item-personality-measure-tipi/"],
        ["ISI/PSQI", "继续隐藏", "未上线", "本地有材料，但PSQI组件脚本存在题序/换算冲突；ISI电子使用授权未完成", "PSQI原始论文可核计分；不能据此替代电子产品授权", "https://www.sleep.pitt.edu/sites/default/files/assets/Instrument%20Materials/PSQI-Article.pdf"],
        ["缺乏父母自主支持问卷", "未查到唯一可靠版本", "未上线", "POPS等候选版本的对象、题数和中文译本不唯一", "需导师确认量表全名、作者、版本和授权", "未查找到可直接上线的唯一版本"],
        ["家庭亲密度与适应性量表", "继续隐藏", "未上线", "FACES II/III/IV及中文修订版题数、维度和阈值不同", "需确认具体版本与中文授权", "未查找到可直接上线的唯一版本"],
    ]
    add_table(doc, ["量表", "结论", "题数", "计分/问题", "来源判断", "公开来源"], public_rows, [1350, 1050, 650, 2350, 2050, 1910])

    for wid, source in [
        ("who5_wellbeing", "WHO官方量表页面；原始分乘4仅用于换算，不能替代专业评估。"),
        ("cognitive_curiosity_student", "南京大学公开中文题页；维度题号与用户给定SPSS计算式一致。"),
        ("big_five_tipi_10", "Gosling作者网站与中文TIPI PDF；短量表维度只作描述性线索。"),
    ]:
        w = worksheet_by_id[wid]
        add_heading(doc, f"1.4.{['who5_wellbeing','cognitive_curiosity_student','big_five_tipi_10'].index(wid)+1} {w['display_title']}", 3)
        add_kv_table(doc, [("题项数量", len(w["questions"])), ("维度", dim_summary(w)), ("计分规则", scoring_summary(w)), ("来源与限制", source)])
        add_table(doc, ["题号", "题项", "选项与分值"], [[q.get("id"), q.get("prompt"), "；".join(f"{o.get('label')}={o.get('score')}" for o in q.get("options", []))] for q in w["questions"]], [1050, 4560, 3750])

    add_heading(doc, "1.5 计分和开放控制", 2)
    add_bullets(doc, [
        "题项、选项和计分均由 content/assessment_worksheets.json 统一驱动，小程序不另存一套题库。",
        "反向计分、维度汇总、区间解释和RFQ-8非线性映射在服务端执行；客户端分值不作为可信输入。",
        "量表开放不等于模型开放。量表可用于自我观察，但聚类模型必须另行通过 artifact hash、准入状态、解释审批和阈值门禁。",
        "高风险文本优先进入人工风险复核，不返回普通画像和训练推荐。",
    ])

    add_page_break(doc)
    add_heading(doc, "二、聚类画像：数据、算法、结果与解释", 1)
    add_heading(doc, "2.1 统一方法", 2)
    add_bullets(doc, [
        "输入是量表维度分或经批准的标准化特征，不以自由文本直接生成稳定人格结论。",
        "离线阶段完成缺失检查、标准化、候选K比较、聚类拟合和画像命名；线上只做与离线同口径的分配。",
        "KMeans模型按到质心的距离分配；GMM模型按高斯混合后验概率分配。具体算法以每个冻结模型工件为准。",
        "PCA只用于二维展示或辅助审查，不代替原始特征空间的正式分配。",
        "模型工件必须通过SHA-256完整性校验；解释文案采用支持性、情境化表述。",
    ])
    model_rows = []
    for m in profiles:
        method = model_method(m)
        model_rows.append([m.get("model_id"), m.get("standard_scale_name"), m.get("worksheet_id") or "未绑定", m.get("n_cases"), m.get("n_features"), method, m.get("chosen_k"), m.get("admission_status")])
    add_table(doc, ["模型ID", "量表/数据", "worksheet", "样本", "特征", "算法", "K", "准入"], model_rows, [1450, 1640, 1450, 600, 600, 1050, 450, 1120])

    add_heading(doc, "2.2 各模型与画像结果", 2)
    for idx, m in enumerate(profiles, 1):
        add_heading(doc, f"2.2.{idx} {m.get('standard_scale_name')}（{m.get('model_id')}）", 3)
        selection = m.get("model_selection", {})
        method = model_method(m)
        add_kv_table(doc, [
            ("来源数据", m.get("source_dataset")),
            ("研究目录", m.get("research_dir")),
            ("样本与特征", feature_summary(m)),
            ("预处理", m.get("preprocessing")),
            ("候选模型与选择", model_selection_summary(selection)),
            ("最终算法/K", f"{method}；K={m.get('chosen_k')}"),
            ("准入/解释审批", f"{m.get('admission_status')} / {m.get('interpretation_approval_status')}"),
            ("完整性", f"artifact_hash={m.get('artifact_hash')}"),
            ("边界", m.get("boundary_notice")),
        ])
        cluster_rows = []
        for c in m.get("clusters", []):
            cluster_rows.append([
                c.get("cluster_id"),
                c.get("display_name") or c.get("profile_name"),
                f"{c.get('n', '—')} / {c.get('percent', '—')}%",
                c.get("product_explanation") or c.get("supportive_explanation"),
                c.get("strength_note"),
                c.get("small_step"),
                c.get("recommended_card_ids"),
            ])
        add_table(doc, ["簇", "画像名", "人数/比例", "支持性描述", "优势线索", "小步骤", "训练卡"], cluster_rows, [500, 1150, 900, 2150, 1400, 1700, 1560])

    add_heading(doc, "2.3 五类必须保留的阻断状态", 2)
    add_table(doc, ["状态", "触发依据", "系统行为", "导师审核重点"], [
        ["low_confidence", "GMM后验概率或距离置信度低于模型阈值", "不输出确定画像；提示结果不稳定，可补充资料或复测", "逐模型冻结阈值并检查校准"],
        ["outlier", "到所有质心距离过大、异常分数或超出训练分布", "阻断自动解释和自动训练推荐，进入人工复核", "确认异常规则、极端值处理和误伤率"],
        ["insufficient_data", "缺题、有效维度不足、样本/记录数未达最低要求", "不硬补、不猜测；说明数据不足", "确认最低有效题数和缺失处理"],
        ["high_risk", "独立风险关键词/结构化风险规则命中", "风险优先；阻断普通画像和普通训练卡，进入人工风险队列", "风险规则不得由聚类或情感得分替代"],
        ["unapproved_model", "模型未获准入、解释未审批或artifact hash不一致", "模型不可用于线上分配；保留内部审计信息", "审批人与版本、工件hash和发布记录"],
    ], [1450, 2500, 3070, 2340])

    add_page_break(doc)
    add_heading(doc, "三、训练卡、量表来源与推荐逻辑", 1)
    add_heading(doc, "3.1 推荐总逻辑", 2)
    add_bullets(doc, [
        "第一道门：风险检查。高风险时不进入普通推荐。",
        "第二道门：量表或画像结果是否完整、模型是否批准、解释是否获准。",
        "第三道门：根据维度阈值或画像簇生成候选训练卡，不以单一总分直接作诊断。",
        "第四道门：过滤不适用场景、禁忌和暂停条件，受控训练卡需要额外批准。",
        "第五道门：候选集通常限制数量，并向用户说明推荐原因，允许在安全候选中共同选择。",
        "训练过程记录最低剂量、完成标准、递进标准和停止规则；不把打卡次数写成疗效证明。",
    ], numbered=True)

    add_heading(doc, "3.2 全部量表—训练卡映射规则", 2)
    rule_rows = []
    for rule in mapping.get("rules", []):
        source = rule.get("source_type")
        condition = rule.get("trigger_condition")
        rule_rows.append([rule.get("rule_id"), source, condition_text(condition), rule.get("recommended_card_ids"), rule.get("reason"), rule.get("approval_status")])
    add_table(doc, ["规则ID", "来源", "触发条件", "候选训练卡", "推荐理由", "审批"], rule_rows, [1500, 1050, 1900, 1800, 2350, 760])

    add_heading(doc, "3.3 34张训练卡逐卡审核", 2)
    reverse_rules: dict[str, list[str]] = {card["id"]: [] for card in cards}
    for rule in mapping.get("rules", []):
        for card_id in rule.get("recommended_card_ids", []):
            reverse_rules.setdefault(card_id, []).append(rule.get("rule_id"))
    for idx, card in enumerate(cards, 1):
        add_heading(doc, f"3.3.{idx} {card.get('title')}（{card.get('id')}）", 3)
        add_kv_table(doc, [
            ("用途", card.get("purpose")),
            ("理论来源", card.get("theory_source")),
            ("目标技能/构念", f"{txt(card.get('target_skill'))}；{txt(card.get('target_constructs'))}"),
            ("基于哪些量表/规则", reverse_rules.get(card["id"]) or "通用候选；没有单独分数触发规则"),
            ("适用", card.get("suitable_for") or card.get("indications")),
            ("不适用/禁忌", card.get("not_suitable_for") or card.get("contraindications")),
            ("时长/最低剂量", f"{txt(card.get('duration_minutes'))}分钟；{txt(card.get('minimum_dose'))}"),
            ("完成标准", card.get("completion_criteria")),
            ("递进标准", card.get("progression_criteria")),
            ("停止规则", card.get("stop_rules")),
            ("循证/安全/发布", f"{txt(card.get('evidence_level'))} / {txt(card.get('safety_level'))} / {txt(card.get('release_policy'))}"),
            ("治理审核", f"{txt(card.get('review_status'))}；{txt(card.get('governance_review_status'))}"),
        ])
        doc.add_paragraph("具体步骤：")
        add_bullets(doc, card.get("steps", []), numbered=True)
        if card.get("example") or card.get("example_phrase"):
            add_note(doc, f"示例：{txt(card.get('example') or card.get('example_phrase'))}")
        if card.get("reflection_questions"):
            doc.add_paragraph("复盘问题：")
            add_bullets(doc, card["reflection_questions"])
        if card.get("fidelity_check"):
            doc.add_paragraph("执行一致性检查：")
            add_bullets(doc, card["fidelity_check"] if isinstance(card["fidelity_check"], list) else [card["fidelity_check"]])

    add_page_break(doc)
    add_heading(doc, "四、三个项目试点方案", 1)
    add_note(doc, "三个项目均为内容协议，不因写入content而自动获得试点资格；须完成导师、伦理、统计和安全审批，并准备中性替代方案。")
    for idx, program in enumerate(programs, 1):
        add_heading(doc, f"4.{idx} {program.get('title')}（{program.get('id')}）", 2)
        add_kv_table(doc, [
            ("目标构念", program.get("target_constructs")),
            ("理论来源", program.get("theory_source")),
            ("对象", program.get("audience")),
            ("状态", f"enabled={program.get('enabled')}；review={program.get('review_status')}；approval={txt(program.get('approval'))}"),
            ("测量计划", program.get("measurement_plan")),
            ("训练卡", program.get("recommended_card_ids")),
            ("最低剂量", program.get("minimum_dose")),
            ("完成定义", program.get("completion_definition")),
            ("暂停标准", program.get("pause_criteria")),
            ("退出标准", program.get("exit_criteria")),
            ("不良反应处理", program.get("adverse_response_plan")),
            ("方案偏离规则", program.get("protocol_deviation_rule")),
            ("纳入标准", program.get("inclusion_criteria")),
            ("排除标准", program.get("exclusion_criteria")),
            ("中性替代", program.get("neutral_alternative")),
            ("结局指标", program.get("outcomes")),
            ("安全门", program.get("safety_gate")),
            ("边界", program.get("boundary_notice")),
        ])
        sessions = program.get("sessions", [])
        if sessions:
            session_rows = []
            for session in sessions:
                if isinstance(session, dict):
                    session_rows.append([session.get("session") or session.get("id") or session.get("week"), session.get("title") or session.get("theme"), session.get("goal") or session.get("content") or session.get("description"), session.get("card_ids") or session.get("activities")])
                else:
                    session_rows.append(["—", txt(session), "—", "—"])
            add_table(doc, ["节次", "主题", "内容/目标", "训练卡/活动"], session_rows, [900, 1800, 3860, 2800])

    add_page_break(doc)
    add_heading(doc, "五、情感计算、语义网络与家庭关系拓扑", 1)
    add_heading(doc, "5.1 基础概念与心理学边界", 2)
    add_bullets(doc, [
        "情感计算是对文本中情绪效价、唤醒程度、表达强度和情绪类别进行结构化估计的技术。本项目不把它当作读心、诊断或危机预测，而是把用户自报文本转换为可复盘的阶段性线索。",
        "社会网络分析用于研究节点及其关系结构。项目中的语义共现网络以概念为节点、同句或同字段共现为边，回答的是“哪些概念经常一起出现”；它不是现实中的朋友关系图，也不能推出真实社会支持质量。",
        "家庭关系拓扑描述经授权家庭绑定中谁与谁存在已确认结构关系、网络是否连通以及数据是否异常。它只描述系统中的结构连接，不测量亲密度、冲突程度、家庭功能或成员好坏。",
        "三类结果都必须与量表、情绪温度计、时间和记录来源共同解释；单次文本、单条边或单个中心性指标不形成稳定结论。",
    ])
    add_heading(doc, "5.2 三条分析线必须分开", 2)
    add_table(doc, ["分析线", "输入", "方法", "当前输出", "不可解释为"], [
        ["情感计算", "经授权的文本记录，按用户/系统/督导来源分层", "分词、词典、否定/程度/重复/转折规则；按记录和时间聚合", "效价、唤醒、强度、类别、覆盖率", "诊断、危机预测或疗效"],
        ["语义共现网络", "人物、场景、想法、情绪、身体、行为和结果概念", "句子/字段级共现；边权与反距离介数；阈值和稳定性", "Top节点/边、反射弧链条、社区（通过门禁时）", "真实社会关系网络"],
        ["家庭关系拓扑审计", "family_links中已确认、未撤回、经授权的结构关系", "HMAC运行级伪名、小单元抑制、连通分量和数据质量", "绑定覆盖率、异常率、聚合结构", "亲密度、支持度或家庭好坏"],
    ], [1450, 2050, 2300, 2060, 1500])
    add_heading(doc, "5.3 项目中的基本处理过程", 2)
    add_bullets(doc, [
        "输入阶段先区分用户自述、系统结构化字段和研究者备注，并检查知情同意、角色权限、撤回状态与最小必要范围。",
        "预处理阶段进行空值、编码、重复记录和时间字段检查；文本分析进行分词、否定词、程度副词、转折和重复强调处理，家庭关系只保留confirmed且未撤回的连接。",
        "情感计算按单条记录生成效价、唤醒、强度和类别，再按用户、时间窗和来源聚合；聚合结果保留覆盖率与样本数，防止少量文本被过度解释。",
        "语义网络从人物、场景、想法、情绪、身体、行为和结果等概念抽取节点，在句子或字段窗口内统计共现，计算加权度、反距离介数、连通结构及稳定性。",
        "家庭拓扑使用运行级HMAC伪名替代真实身份，计算节点、连接、连通分量、孤立点、重复边和异常关系，仅输出达到小单元阈值的聚合结构。",
        "输出阶段写入不含原始文本的聚合JSON，研究后台通过只读接口展示；普通用户端不直接展示未经验证的实时算法判断。",
    ], numbered=True)
    add_heading(doc, "5.4 当前实现与治理状态", 2)
    add_bullets(doc, [
        "主路径是 analysis/text_analysis/；后端通过 GET /api/text-analysis/summary 只读聚合JSON，仅管理员和研究者可访问。",
        "分析脚本使用只读输入和运行清单，最终产物不保留原始自由文本；运行ID和伪名采用密钥化方式，密钥不入仓。",
        "语义网络介数中心性需使用 distance=1/(weight+epsilon)，避免把高共现错误当作远距离。",
        "家庭拓扑只纳入 confirmed 且未撤回关系；小单元默认抑制，低于阈值不展示稀有节点或边。",
        "现有词典仍是小规模项目词典，真实数据记录数、覆盖率、人工标注F1和外部效标尚不足以支持用户端解释。",
        "高风险仍由独立风险规则与人工队列处理，情感得分和网络中心性不得替代。",
    ])
    add_heading(doc, "5.5 从初版到当前版本的优化过程", 2)
    add_bullets(doc, [
        "从单一情绪词命中升级为效价、唤醒、强度和类别的分层输出，并加入否定、程度、重复与转折规则，降低机械关键词计数造成的误判。",
        "从直接保留文本结果升级为运行清单、输入哈希、质量状态和无原文聚合产物，便于复现、比较和隐私审计。",
        "语义网络修正了加权边的距离定义：以1/(weight+epsilon)作为最短路距离，避免高共现边被错误解释为更远。",
        "家庭拓扑增加confirmed/撤回过滤、运行级伪名、小单元抑制和异常边审计，使结构分析与业务授权状态一致。",
        "API从任意前端读取升级为研究者/管理员只读、统一响应契约与审计留痕；质量状态不是valid时不向前端提供解释。",
        "高风险识别继续独立于情感计算和网络指标，避免把低效价、高中心性或异常结构误当作危机证据。",
    ])
    add_heading(doc, "5.6 后续准入前必须补齐", 2)
    add_bullets(doc, [
        "经授权中文词典、版本和许可证清单；",
        "双人独立标注、一致性、各类别precision/recall/F1与macro-F1；",
        "与情绪温度计愉悦度、唤醒度、强度的外部效标验证；",
        "语义网络的边阈值、bootstrap稳定性和社区稳定性；",
        "家庭拓扑真实数据可用性、撤回过滤和小单元隐私验收；",
        "产物valid/empty/insufficient_data/stale/validation_failed/privacy_blocked六态。",
    ])

    add_page_break(doc)
    add_heading(doc, "六、账户、鉴权、角色和研究者策略", 1)
    add_heading(doc, "6.1 角色分类使用", 2)
    add_table(doc, ["角色", "主要入口", "允许操作", "限制"], [
        ["家长端", "微信小程序", "目标、亲子情绪事件、家长量表、训练卡、打卡、周报和督导请求", "只能访问本人及经授权家庭范围；不见研究分析原始数据"],
        ["学生端", "小程序学生支持性测评入口", "学生量表、阶段性画像、学生训练卡和复测", "未成年人保护；不展示诊断标签、复杂网络或未经批准模型"],
        ["研究者", "Web研究后台/只读接口", "脱敏汇总、模型审计、聚合情感/网络结果和受控导出", "最小权限；不能查看不必要原文；高风险导出需二次确认并审计"],
        ["督导/管理员", "Web后台", "风险复核、人工备注、内容与模型审批、账户治理", "操作留痕；不得用系统结论替代专业判断"],
    ], [1200, 1800, 3300, 3060])
    add_heading(doc, "6.2 鉴权与研究者账号策略", 2)
    add_bullets(doc, [
        "登录支持微信身份链路及受控账号方式；生产环境不得依赖默认demo用户或默认管理员token。",
        "服务端以签名token和数据库角色为准，不信任客户端自行声明的user_id或role。",
        "研究者账号由管理员创建，按项目和数据范围最小授权；禁止多人共享账号。",
        "建议启用强密码、定期轮换、到期停用、人员离组即时撤权和异常访问复核。",
        "查看画像详情、导出、风险记录和人工复核均写入audit_logs；高风险或需复核数据导出需要显式确认。",
        "研究者默认只见匿名ID、聚合统计和必要字段；联系方式、微信标识、原文和家庭逐边关系不默认开放。",
    ])

    add_heading(doc, "七、数据存放、数据库表与完整数据流", 1)
    add_heading(doc, "7.1 数据存放位置", 2)
    add_table(doc, ["环境", "位置", "内容", "治理"], [
        ["内容库", "仓库 content/*.json", "量表、训练卡、映射规则、项目方案、模型工件", "版本控制；内容审核与模型审批分离"],
        ["本地开发", "backend 本地SQLite（具体路径由环境配置）", "联调和测试数据", "不得提交*.db/*.sqlite3；不能当生产研究库"],
        ["生产", "CloudBase云托管所连接的受控数据库/MySQL", "账户、测评、画像、风险、审计等业务数据", "凭证不入仓；备份、访问控制、保留与删除按伦理方案"],
        ["离线研究", "经授权脱敏快照与 outputs 下聚合产物", "画像训练数据、情感/网络聚合", "原始快照隔离；hash、manifest、小单元抑制；产物不含原文"],
    ], [1250, 2450, 2900, 2760])
    add_heading(doc, "7.2 主要数据库表", 2)
    table_rows = [
        ["users", "账户、角色、状态和登录身份"], ["family_links", "家长—学生等经授权家庭绑定"],
        ["goals", "用户目标"], ["emotion_diaries / emotion_thermometer", "情绪事件与温度计结构化记录"],
        ["assessment_worksheets", "可执行测评版本快照"], ["assessment_results", "服务端计分后的测评结果"],
        ["student_profiles / student_profile_followups", "学生画像结果与后续复测"], ["parent_assessment_submissions / parent_report_actions", "家长量表提交与报告动作"],
        ["profile_reviews", "画像人工复核"], ["risk_review_records", "高风险命中、处置和人工队列"],
        ["training_cards / checkins", "训练卡与练习打卡"], ["weekly_reports", "周度聚合报告"],
        ["supervision_requests / messages", "督导请求和消息"], ["records", "统一研究摘要索引"],
        ["consent_records / privacy_requests", "知情同意、撤回、查阅和删除请求"], ["audit_logs", "敏感查看、导出、复核和治理操作日志"],
        ["relationship_* 系列表", "家庭关系试点纳入、筛查、任务、叙事、纵向、假设反馈与研究备注"],
    ]
    add_table(doc, ["表/表组", "用途"], table_rows, [3300, 6060])

    add_heading(doc, "7.3 一名用户从登录到收到结果的全过程示例", 2)
    add_bullets(doc, [
        "用户在小程序发起登录；服务端核验微信身份或受控账号凭证，在users中定位/创建账户并返回短期访问token。",
        "小程序携带token请求量表列表；后端只返回enabled_for_user=true且内容版本可用的worksheet。",
        "用户填写题项。客户端只保存题号和所选值，不把客户端计算的总分当成可信结果。",
        "提交到测评接口后，服务端校验未知题号、重复题、非法选项、缺失必答题和worksheet版本；不合格请求直接拒绝。",
        "服务端按冻结规则进行正向/反向、维度、总分、RFQ-8映射等计算，并先执行独立风险检查。",
        "若命中high risk：普通画像和训练建议被阻断，写入risk_review_records并提示人工支持路径。",
        "若风险通过：在同一受控流程中写assessment_results；需要画像时再检查模型准入、解释审批、artifact hash、数据完整性、离群和置信度。",
        "通过画像门禁后写student_profiles/records；未通过则保存事实分数和明确阻断状态，不强行贴画像。",
        "推荐引擎根据量表维度或批准画像生成候选训练卡，过滤禁忌、高风险和未批准卡，返回推荐理由和边界说明。",
        "小程序收到结构化结果并展示支持性解释、可选择的小步骤和非诊断声明；后续打卡、复测、周报与人工复核分别写对应表并留审计。",
    ], numbered=True)

    add_heading(doc, "八、质量控制、当前限制与导师待审事项", 1)
    add_heading(doc, "8.1 本轮已实现并验证的内容", 2)
    add_bullets(doc, [
        "项目负责人批准后，33份已有完整worksheet的量表进入试点开放；仅目录元数据且缺少执行链的量表继续阻断。",
        "RFQ-8新增非线性映射均值计分支持；服务端仍统一校验和计分。",
        "新增量表—训练卡规则，并保留TIPI仅作描述、不因人格维度自动推荐的限制。",
        "HPLP统一采用牛至旭40题六维版本并绑定对应主模型；高鸣聪HPLP模型保留为历史对照。",
        "全量量表审计达到33份、0 blocker、0 error、0 warning；内容校验与专项回归通过。",
    ])
    add_heading(doc, "8.2 仍需导师、伦理或统计负责人确认", 2)
    add_bullets(doc, [
        "每份量表的正式中文版本、电子使用授权、对象年龄、反向题、维度题号和解释阈值；",
        "项目负责人批准已记录；仍需归档各量表正式中文版本、电子使用范围及必要的第三方权利证明；",
        "聚类画像的训练样本代表性、K选择、阈值校准、簇稳定性、外部效标和命名；",
        "训练卡的理论对应、最低剂量、停止规则、危机转介和受控发布卡清单；",
        "三个项目试点的纳排标准、样本量、主要/次要结局、对照条件和不良事件流程；",
        "研究者可见字段、数据保留期、撤回/删除流程、小单元阈值以及跨端角色边界。",
    ])
    add_heading(doc, "8.3 建议纳入导师汇报的补充内容", 2)
    add_bullets(doc, [
        "把“量表开放”和“模型准入”作为两张独立签字表，避免题库可用被误解为聚类有效。",
        "汇报模型时同时展示失败状态和不解释条件，而不仅展示簇名称。",
        "汇报训练卡时展示从量表/画像到候选集、共享选择、剂量、复盘和停止的完整链路。",
        "将情感计算、语义网络、家庭拓扑分别汇报，明确当前均属于离线研究工具，不进入普通用户实时判断。",
        "把版本、授权、伦理、统计、隐私、模型工件hash和审计日志纳入每轮试点的发布清单。",
    ])

    add_page_break(doc)
    add_heading(doc, "附录A：全部worksheet逐题、维度与计分清单", 1)
    add_note(doc, "本附录直接从当前content/assessment_worksheets.json生成，是小程序和服务端共用的当前内容快照。每份量表列出题项、选项、维度、计分和审批状态；仅目录元数据而无worksheet的项目不在本附录伪造题项。")
    for idx, w in enumerate(worksheets, 1):
        add_heading(doc, f"A.{idx} {w.get('display_title')}（{w.get('id')}）", 2)
        add_kv_table(doc, [
            ("对象/分类", f"{txt(w.get('audience'))}；{txt(w.get('audience_class'))}；{txt(w.get('category'))}"),
            ("状态", f"enabled_for_user={w.get('enabled_for_user')}；review={txt(w.get('review_status'))}"),
            ("题项数量", len(w.get("questions", []))),
            ("维度", dim_summary(w)),
            ("计分", scoring_summary(w)),
            ("推荐训练卡", w.get("recommended_card_ids")),
            ("来源", f"{txt(w.get('source_title'))}；{txt(w.get('source_file'))}；{txt(w.get('source_version'))}"),
            ("审核备注", w.get("review_note")),
        ])
        add_table(doc, ["题号", "题项", "选项与分值"], [[q.get("id"), q.get("prompt"), "；".join(f"{o.get('label')}={o.get('score')}" for o in q.get("options", []))] for q in w.get("questions", [])], [1050, 4610, 3700])

    add_page_break(doc)
    add_heading(doc, "附录B：课程内容库补充", 1)
    add_note(doc, "课程不是量表，也不是聚类模型；它们用于把训练卡组织成可学习、可练习、可复盘的短内容。")
    for idx, course in enumerate(courses, 1):
        add_heading(doc, f"B.{idx} {course.get('title')}（{course.get('id')}）", 2)
        add_kv_table(doc, [
            ("主题/场景", f"{txt(course.get('theme'))}；{txt(course.get('scene'))}"),
            ("时长", f"{txt(course.get('duration_minutes'))}分钟"),
            ("学习目标", course.get("learning_objectives")),
            ("核心概念", course.get("core_concept")),
            ("常见误区", course.get("common_misconceptions")),
            ("正例", course.get("worked_example")),
            ("反例", course.get("counter_example")),
            ("引导练习", course.get("guided_practice")),
            ("迁移任务", course.get("transfer_task")),
            ("复盘", course.get("reflection_prompts")),
            ("加强计划", course.get("booster_plan")),
            ("关联训练卡/项目", course.get("relation_to_cards_or_programs")),
            ("边界", course.get("boundary_notice")),
        ])

    add_page_break(doc)
    add_heading(doc, "附录C：审核证据、公开来源与项目路径", 1)
    add_table(doc, ["类别", "路径/来源", "用途"], [
        ["量表执行", "content/assessment_worksheets.json", "题项、选项、计分、维度、开放状态"],
        ["量表目录", "content/scales_catalog.json", "来源、审核和开放登记"],
        ["训练卡", "content/training_cards.json", "34张训练卡完整内容"],
        ["推荐规则", "content/assessment_training_map.json", "量表/画像到训练卡候选集"],
        ["聚类模型", "content/profiles/profile_*.json", "11个冻结模型工件"],
        ["试点", "content/programs.json", "三个项目方案"],
        ["课程", "content/courses.json", "五门课程内容"],
        ["情感/SNA", "docs/00_当前事实基准/情感计算与社会网络分析离线原型优化与实现任务路径_20260711.md", "方法、问题、治理和验收路径"],
        ["公开来源", "WHO、南京大学、Gosling作者页、Pittsburgh PSQI原始文献、PeerJ RFQ-8评分表等", "核对题项、计分、维度和使用边界"],
    ], [1600, 4850, 2910])
    add_note(doc, "版本结论：本材料按2026-07-12工作区内容生成。项目负责人试点批准已记录；第三方版权证据、模型外部效度和独立伦理审批仍以最终归档文件为准。")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
